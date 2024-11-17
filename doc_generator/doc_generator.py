from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from mimesis import Text
from docx.enum.style import WD_STYLE_TYPE
import math2docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import io
import random
from faker import Faker
import pandas as pd


from mimesis import Generic
from mimesis.locales import Locale
from mimesis.random import Random

# Faker будем использовать и для английского и для русского, т.к. в примере документа оба языка
fake = Faker(['en_US', 'ru_RU'])
text_mimesis = Text('ru')
NUM_ITERATIONS = 5  # Количество итераций для генерации содержимого

COLORS = """000000 000080 00008B 0000CD 0000FF 006400 008000 008080 008B8B 00BFFF 00CED1 
00FA9A 00FF00 00FF7F 00FFFF 191970 1E90FF 20B2AA 228B22 2E8B57 2F4F4F 32CD32 
3CB371 40E0E0 4169E1 4682B4 483D8B 48D1CC 4B0082 556B2F 5F9EA0 6495ED 66CDAA 
696969 6A5ACD 6B8E23 708090 778899 7B68EE 7CFC00 7FFF00 7FFFD4 800000 800080 
808000 808080 87CEEB 87CEFA 8A2BE2 8B0000 8B008B 8B4513 8FBC8F 90EE90 9370D8 
9400D3 98FB98 9932CC 9ACD32 A0522D A52A2A A9A9A9 ADD8E6 ADFF2F AFEEEE B0C4DE 
B0E0E6 B22222 B8860B BA55D3 BC8F8F BDB76B C0C0C0 C71585 CD5C5C CD853F D2691E 
D2B48C D3D3D3 D87093 D8BFD8 DA70D6 DAA520 DC143C DCDCDC DDA0DD DEB887 E0FFFF 
E6E6FA E9967A EE82EE EEE8AA F08080 F0E68C F0F8FF F0FFF0 F0FFFF F4A460 F5DEB3 
F5F5DC F5F5F5 F5FFFA F8F8FF FA8072 FAEBD7 FAF0E6 FAFAD2 FDF5E6 FF0000 FF00FF 
FF00FF FF1493 FF4500 FF6347 FF69B4 FF7F50 FF8C00 FFA07A FFA500 FFB6C1 FFC0CB 
FFD700 FFDAB9 FFDEAD FFE4B5 FFE4C4 FFE4E1 FFEBCD FFEFD5 FFF0F5 FFF5EE FFF8DC 
FFFACD FFFAF0 FFFAFA FFFF00 FFFFE0 FFFFF0 FFFFFF"""


def change_orientation(document, landscape=True):
    """ Функция смены ориентации листа"""
    section = document.sections[-1]
    if landscape and section.orientation == WD_ORIENT.PORTRAIT:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    elif not landscape and section.orientation == WD_ORIENT.LANDSCAPE:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width, section.page_height = section.page_height, section.page_width


def add_header(document, base_font_size):
    section = document.sections[-1]
    header = section.header
    header_text = fake.sentence(nb_words=5)

    header_paragraph = header.paragraphs[0]
    header_paragraph.text = header_text
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_paragraph.runs[0]
    run.font.size = Pt(base_font_size - 4)


def add_footer(document, base_font_size):
    section = document.sections[-1]
    footer = section.footer
    footer_text = fake.sentence(nb_words=5)

    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = footer_text
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_paragraph.runs[0]
    run.font.size = Pt(base_font_size - 4)


class RandomHeadingFormat:
    def __init__(self, random_paragraph_format):
        self.font_size_1 = random_paragraph_format.font_size + random.choice([2, 4, 6, 8])
        self.font_size_2 = random_paragraph_format.font_size + random.choice([2, 4, 6])
        self.font_size_3 = random_paragraph_format.font_size + random.choice([2, 4])
        self.bold = random.random() < 0.7
        self.italic = random.random() < 0.3
        self.alignment = random.choices([WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT],
                                        weights=[30, 70, 30])[0]


def add_heading(document, random_heading_format):
    """
    Функция с добавлением заголовка
    """
    # Генерация заголовка
    level = random.randint(0, 3)  # Уровень заголовка
    if level == 1:
        font_size = random_heading_format.font_size_1
    elif level == 2:
        font_size = random_heading_format.font_size_2
    else:
        font_size = random_heading_format.font_size_3

    heading_text = fake.sentence(nb_words=random.randint(3, 7))  # Рандомная генерация фейкером / мимезисом
    heading = document.add_heading(level=level)

    # Добавление нумерации заголовка с вероятностью 50%
    if random.random() < 0.5:
        numbering = f"{random.randint(1, 15)}. "
        heading_text = numbering + heading_text
    else:
        heading_text = heading_text

    run = heading.add_run(heading_text)
    run.bold = random.choice([True, False])
    if not run.bold:
        run.italic = True
    else:
        run.italic = random.choice([True, False])
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Выравнивание заголовка (по условию задачи чаще центр)
    heading.paragraph_format.alignment = random_heading_format.alignment


class RandomParagraphFormat:
    def __init__(self):
        self.left_indent = Inches(0.25) if random.choice([True, False]) else Inches(0)
        self.space_before = random.randint(3, 10)
        self.space_after = random.randint(3, 10)
        self.first_line_indent = Inches(0.25) if random.choice([True, False]) else Inches(0)
        self.alignment = random.choice(
            [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.JUSTIFY])
        self.font_size = random.randint(10, 16)
        self.font_name = random.choice(
            ['Times New Roman', 'Arial', 'Calibri', 'Georgia', 'Verdana', 'Tahoma', 'Garamond', 'Helvetica',
             'Courier New', 'Trebuchet MS', 'Comic Sans'])
        self.line_spacing = random.uniform(1.0, 1.4)


def add_paragraph(document, random_paragraph_format):
    """ Функция генерации обычного текста """
    # Генерация абзаца
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = random_paragraph_format.left_indent  # Рандомная красная строка
    paragraph_format.space_before = Pt(random_paragraph_format.space_before)
    paragraph_format.space_after = Pt(random_paragraph_format.space_after)
    paragraph_format.first_line_indent = random_paragraph_format.first_line_indent

    # Выравнивание абзаца
    paragraph.alignment = random_paragraph_format.alignment
    
    paragraph_format.line_spacing = random_paragraph_format.line_spacing

    # Генерация текста для абзаца
    if random.choice([True, False]):
        text = fake.text(max_nb_chars=400)
    else:
        text = text_mimesis.text(quantity=random.randint(5, 10))
    run = paragraph.add_run(text)
    run.font.size = Pt(random_paragraph_format.font_size)
    run.font.name = random_paragraph_format.font_name


def get_table(doc, base_font_size):
    random = Random()
    generic = Generic(locale=Locale.RU)

    below_above = random.randint(0, 1)
    caption_text = fake.sentence(nb_words=random.randint(3, 7))
    if random.choice([True, False]):
        caption_text = f"Табл. {random.randint(1, 100)} - {caption_text}"
    elif random.choice([True, False]):
        caption_text = f"Таблица {random.randint(1, 100)} - {caption_text}"
    else:
        caption_text = f"Таблица. {caption_text}"
    if below_above == 1:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(caption_text)
        run.font.size = Pt(base_font_size)

    rows = random.randint(a=3, b=7)
    cols = random.randint(a=3, b=7)
    table = doc.add_table(rows=0, cols=cols)
    table.style = random.choice_enum_item(('Light Grid',
                                           'Light Grid Accent 1',
                                           'Light Grid Accent 2',
                                           'Light Grid Accent 3',
                                           'Light Grid Accent 4',
                                           'Light Grid Accent 5',
                                           'Light Grid Accent 6',
                                           'Medium Grid 1',
                                           'Medium Grid 1 Accent 1',
                                           'Medium Grid 1 Accent 2',
                                           'Medium Grid 1 Accent 3',
                                           'Medium Grid 1 Accent 4',
                                           'Medium Grid 1 Accent 5',
                                           'Medium Grid 1 Accent 6',
                                           'Medium Grid 2',
                                           'Medium Grid 2 Accent 1',
                                           'Medium Grid 2 Accent 2',
                                           'Medium Grid 2 Accent 3',
                                           'Medium Grid 2 Accent 4',
                                           'Medium Grid 2 Accent 5',
                                           'Medium Grid 2 Accent 6',
                                           'Table Grid'))
    for row in range(rows):
        row_cells = table.add_row().cells
        for col in range(cols):
            row_cells[col].text = ' '.join(generic.text.words(quantity=random.randint(a=1, b=3)))

    if below_above == 0:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(caption_text)
        run.font.size = Pt(base_font_size)


def generate_plot_or_chart():
    chart_type = random.choice(['line', 'bar', 'scatter', 'pie'])
    fig, ax = plt.subplots()
    if chart_type == 'line':
        x = range(10)
        y = [random.randint(1, 100) for _ in x]
        ax.plot(x, y, color=random.choice(['red', 'green', 'blue']), linestyle=random.choice(['-', '--', '-.', ':']),
                linewidth=random.uniform(0.5, 2.5))
        ax.set_title(fake.sentence(nb_words=random.randint(1, 4)))
        ax.set_xlabel(fake.word())
        ax.set_ylabel(fake.word())

    elif chart_type == 'bar':
        x = range(5)
        y = [random.randint(1, 100) for _ in x]
        ax.bar(x, y, color=random.choice(['red', 'green', 'blue']))
        ax.set_title(fake.sentence(nb_words=random.randint(1, 4)))
        ax.set_xlabel(fake.word())
        ax.set_ylabel(fake.word())

    elif chart_type == 'scatter':
        x = [random.uniform(0, 100) for _ in range(20)]
        y = [random.uniform(0, 100) for _ in range(20)]
        ax.scatter(x, y, color=random.choice(['red', 'green', 'blue']), s=random.randint(10, 200))
        ax.set_title(fake.sentence(nb_words=random.randint(1, 4)))
        ax.set_xlabel(fake.word())
        ax.set_ylabel(fake.word())

    elif chart_type == 'pie':
        sizes = [random.randint(1, 10) for _ in range(4)]
        labels = [fake.word(), fake.word(), fake.word(), fake.word()]
        colors = ['red', 'green', 'blue', 'yellow']
        ax.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=140)
        ax.set_title(fake.sentence(nb_words=random.randint(1, 4)))


def add_picture_with_caption(document, base_font_size):
    """ Функция добавления картинки вместе с подписью сразу """
    # Создание графика (будет вместо рисунков)
    generate_plot_or_chart()
    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='PNG')
    plt.close()
    img_stream.seek(0)

    # Определяем выравнивание
    alignment = random.choices(['left', 'center', 'right'], weights=[30, 70, 30])[0]
    alignment_value = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT
    }[alignment]

    # Создание нового параграфа для рисунка
    paragraph = document.add_paragraph()  # Новый параграф для рисунка
    run = paragraph.add_run()  # Добавляем пустой run для размещения изображения

    # Добавление рисунка
    run.add_picture(img_stream, width=Inches(4))

    # Устанавливаем выравнивание для рисунка
    paragraph.alignment = alignment_value

    # Определение позиции подписи (80% - снизу, 20% - сверху)
    position = random.choices(['below', 'above'], weights=[0.8, 0.2])[0]

    # Создание подписи
    caption_prefix = random.choice(["Рис.", "Рисунок", "Figure"])
    picture_number = random.randint(1, 100)
    caption_text = f"{caption_prefix} {picture_number}. {fake.sentence(nb_words=random.randint(3, 7))}"

    # Создаем подпись в новом параграфе
    caption_paragraph = document.add_paragraph(caption_text)  # Добавляем подпись
    caption_paragraph.alignment = alignment_value  # Устанавливаем выравнивание для подписи

    # Устанавливаем стиль для подписи
    caption_paragraph.style = 'Caption'
    run = caption_paragraph.runs[0]  # Получаем доступ к тексту подписи
    run.font.size = Pt(max(base_font_size - 2, 8))
    run.italic = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Если подпись выше рисунка, нужно вставить перевод строки перед подписью
    if position == 'above':
        document.add_paragraph()  # Добавляем пустой параграф для отделения

    # Убедимся, что у нас есть отступ между рисунком и подписью
    document.add_paragraph()  # Добавляем пустой параграф для отделения


def get_formula(doc, latex_data):
    """ Вызвать функцию == добавить формулу в документ"""
    paragraph = doc.add_paragraph()
    latex_ = r'{}'.format(latex_data[random.randint(a=0, b=2371)])

    math2docx.add_math(paragraph, latex_)


def add_numbered_list(document, random_paragraph_format):
    """ Вызвать функцию == добавить нумерованный список в документ"""
    list_type = 'List Number'
    num_items = random.randint(3, 7)
    for _ in range(num_items):
        list_item = fake.sentence(nb_words=random.randint(3, 15)) if random.choice([True, False]) else \
            text_mimesis.text(quantity=1).split('.')[0]
        paragraph = document.add_paragraph(style=list_type)
        run = paragraph.add_run(list_item)
        run.font.size = Pt(random_paragraph_format.font_size)
        run.font.name = random_paragraph_format.font_name
        paragraph.paragraph_format.left_indent = Inches(0.85)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = random_paragraph_format.line_spacing - 0.2

    empty_paragraph = document.add_paragraph(style="Normal")
    empty_paragraph.paragraph_format.space_before = Pt(0)
    empty_paragraph.paragraph_format.space_after = Pt(0)
    empty_paragraph.paragraph_format.line_spacing = Pt(2)
    empty_run = empty_paragraph.add_run(' ')
    empty_run.font.size = Pt(1)


def add_bulleted_list(document, random_paragraph_format):
    """ Вызвать функцию == добавить маркированный список в документ"""
    list_type = 'List Bullet'
    num_items = random.randint(3, 7)
    for _ in range(num_items):
        list_item = fake.sentence(nb_words=random.randint(3, 15)) if random.choice([True, False]) else \
            text_mimesis.text(quantity=1).split('.')[0]
        paragraph = document.add_paragraph(style=list_type)
        run = paragraph.add_run(list_item)
        run.font.size = Pt(random_paragraph_format.font_size)
        run.font.name = random_paragraph_format.font_name
        paragraph.paragraph_format.left_indent = Inches(0.85)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = random_paragraph_format.line_spacing - 0.2
    
    empty_paragraph = document.add_paragraph(style="Normal")
    empty_paragraph.paragraph_format.space_before = Pt(0)
    empty_paragraph.paragraph_format.space_after = Pt(0)
    empty_paragraph.paragraph_format.line_spacing = Pt(1)
    empty_run = empty_paragraph.add_run(' ')
    empty_run.font.size = Pt(1)


def add_footnotes_section(document, random_paragraph_format):
    document.add_paragraph()
    for fn_count in range(random.randint(1, 10)):
        fn_paragraph = document.add_paragraph()
        index_run = fn_paragraph.add_run(f'[{fn_count}] ')
        index_run.bold = True
        fn_paragraph.add_run(fake.sentence(nb_words=random.randint(3, 15)))
        fn_paragraph.runs[0].font.size = Pt(random_paragraph_format.font_size)
        fn_paragraph.runs[0].font.name = random_paragraph_format.font_name
        fn_paragraph.runs[1].font.size = Pt(random_paragraph_format.font_size)
        fn_paragraph.runs[1].font.name = random_paragraph_format.font_name
        fn_paragraph.style = document.styles['Normal']
        fn_paragraph.paragraph_format.left_indent = Pt(18)
        fn_paragraph.paragraph_format.space_after = Pt(2)
        fn_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fn_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_multicolumn_text(doc, based_font_size):
    generic = Generic(locale=Locale.RU)
    random = Random()

    colums = random.weighted_choice({'2': 0.5, '3': 0.5})
    if colums == '2':
        table = doc.add_table(rows=1, cols=2)

        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)

        quantity = random.randint(a=30, b=60)
        left_cell.text = ' '.join(generic.text.words(quantity=quantity))
        left_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_cell.paragraphs[0].runs[0].font.size = Pt(based_font_size)
        right_cell.text = ' '.join(generic.text.words(quantity=quantity))
        right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        right_cell.paragraphs[0].runs[0].font.size = Pt(based_font_size)

    elif colums == '3':
        table = doc.add_table(rows=1, cols=3)

        left_cell = table.cell(0, 0)
        center_cell = table.cell(0, 1)
        right_cell = table.cell(0, 2)

        quantity = random.randint(a=20, b=40)
        left_cell.text = ' '.join(generic.text.words(quantity=quantity))
        left_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_cell.paragraphs[0].runs[0].font.size = Pt(based_font_size)
        right_cell.text = ' '.join(generic.text.words(quantity=quantity))
        right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        right_cell.paragraphs[0].runs[0].font.size = Pt(based_font_size)
        center_cell.text = ' '.join(generic.text.words(quantity=quantity))
        center_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        center_cell.paragraphs[0].runs[0].font.size = Pt(based_font_size)


def get_footnote(doc, based_font_size):
    generic = Generic(locale=Locale.RU)
    random = Random()

    section = doc.sections[0]
    paragraph = section.footer.paragraphs[0]
    paragraph.style.font.size = Pt(based_font_size)

    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    power_utf8 = {1: '\u00B9', 2: '\u00B2', 3: '\u00B3', 4: '\u2074'}
    for note in range(1, random.randint(a=2, b=5)):
        random_word = generic.text.words(quantity=1)
        random_sentence = generic.text.words(quantity=random.randint(a=3, b=6))
        if random.choice([True, False]):
            footnote_text = f"{power_utf8[note]}{random_word[0]} — {' '.join(random_sentence)} \n"
        else:
            footnote_text = f"{random_word[0]}{power_utf8[note]} — {' '.join(random_sentence)} \n"
        paragraph.add_run(footnote_text).italic = True


def get_image(doc, based_font_size, num_of_pictures):
    random = Random()
    generic = Generic(locale=Locale.RU)

    image_note = random.randint(a=0, b=1)
    if image_note == 0:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(str('Рисунок ' + str(num_of_pictures) + ' — ' + ' '.join(
            generic.text.words(quantity=random.randint(a=1, b=3)))))
        run.font.size = Pt(based_font_size)
        num_of_pictures += 1
    if random.weighted_choice({0: 0.5, 1: 0.5}) > 0:
        doc.add_picture(
            'doc_generator/Caltech101/' + str(random.randint(a=1, b=9142)) + '.jpg',
            width=Inches(random.choice_enum_item((3.5, 4, 4.5))))
    else:
        doc.add_picture('doc_generator/PlotQA/' + str(random.randint(a=0, b=33656)) + '.png',
                        width=Inches(random.choice_enum_item((5, 5.5, 6))))
    if image_note == 1:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(str('Рисунок ' + str(num_of_pictures) + ' — ' + ' '.join(
            generic.text.words(quantity=random.randint(a=1, b=3)))))
        run.font.size = Pt(based_font_size)
        num_of_pictures += 1


def generate_document(path):
    """
    Запуск генерации документа.
    Создает документ, выбирает базовый размер шрифта и сам шрифт.
    После этого выбирает количество итераций генерации элементов.
    Генерирует NUM_ITERATIONS раз все элементы, вызывая функции выше.
    """
    document = Document()
    table_styles = [style.name for style in document.styles if style.type == WD_STYLE_TYPE.TABLE]

    # Начальные характеристики документа
    random_paragraph_format = RandomParagraphFormat()
    random_heading_format = RandomHeadingFormat(random_paragraph_format)

    num_of_pictures = 1
    latex_data = pd.read_csv(r'doc_generator/latex_for_formulas.csv', index_col=False)['formula']

    footnote_type = random.choices([0, 1], weights=[40, 60])[0]

    element_funcs = []
    element_funcs += [lambda: add_heading(document, random_heading_format)] * random.randint(1, 3)
    element_funcs += [lambda: add_paragraph(document, random_paragraph_format)] * random.randint(1, 5)
    element_funcs += [lambda: get_table(document, random_paragraph_format.font_size)] * random.randint(1, 2)
    element_funcs += [lambda: add_picture_with_caption(document, random_paragraph_format.font_size)] * random.randint(1, 2)
    element_funcs += [lambda: add_numbered_list(document, random_paragraph_format)] * random.randint(1, 3)
    element_funcs += [lambda: add_bulleted_list(document, random_paragraph_format)] * random.randint(1, 3)
    element_funcs += [lambda: get_formula(document, latex_data)] * random.randint(1, 3)
    element_funcs += [lambda: add_multicolumn_text(document, random_paragraph_format.font_size)] * random.randint(1, 3)

    for i in range(NUM_ITERATIONS):
        # Добавление новой секции на новой странице, кроме первой
        # А на первой добавляем колонтитулы -- раз и на все страницы.
        if i != 0:
            document.add_section(WD_SECTION.NEW_PAGE)
        elif (footnote_type == 1):
            add_header(document, random_paragraph_format.font_size)
            get_footnote(document, random_paragraph_format.font_size)
        else:
            add_header(document, random_paragraph_format.font_size)
            add_footer(document, random_paragraph_format.font_size)

        # Изменение ориентации страницы на вертикальную с вероятностью 85%
        if random.random() < 0.85:
            landscape = random.choice([True, False])
            change_orientation(document, landscape=landscape)

        # Добавление заголовка
        add_heading(document, random_heading_format)

        # Добавление абзаца
        add_paragraph(document, random_paragraph_format)

        # Для фикса бага с мультиколонками
        last_was_heading = False

        random.shuffle(element_funcs)
        for func in element_funcs:
            # Скипаем мультиколонку если прошлый элемент заголовок
            if last_was_heading and 'add_multicolumn_text' in func.__code__.co_names:
                continue
            if 'add_heading' in func.__code__.co_names:
                last_was_heading = True
            else:
                last_was_heading = False
            func()

        if footnote_type == 0:
            add_footnotes_section(document, random_paragraph_format)

    document.save(path)